# -*- coding: utf-8 -*-
"""Generate BAO-CAO-MEDICARE-STORE.docx"""
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:\thietbiyte_shop\BAO-CAO-MEDICARE-STORE.docx"

doc = Document()

# Styles helper
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    if bold:
        run.bold = True

def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

# ===== CONTENT =====
add_title("BÁO CÁO DỰ ÁN")
p = doc.add_paragraph()
r = p.add_run("MEDICARE STORE – WEBSITE THƯƠNG MẠI ĐIỆN TỬ THIẾT BỊ Y TẾ")
r.bold = True
r.font.size = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Fullstack: React + Node.js + SQL Server").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

add_h1("1. GIỚI THIỆU DỰ ÁN")
add_p(
    "MediCare Store là hệ thống website thương mại điện tử chuyên bán thiết bị y tế "
    "(máy đo huyết áp, máy đo đường huyết, khẩu trang, thiết bị oxy, nhiệt kế, xe lăn…). "
    "Dự án gồm giao diện khách hàng (storefront), trang quản trị (admin dashboard) và "
    "REST API backend kết nối cơ sở dữ liệu Microsoft SQL Server."
)
add_p("Mục tiêu: cung cấp trải nghiệm mua sắm trực tuyến hiện đại, quản lý sản phẩm/đơn hàng "
      "tập trung, hỗ trợ thanh toán COD và online (mô phỏng), đánh giá sản phẩm có ảnh.")

add_h1("2. CÔNG NGHỆ SỬ DỤNG")
add_h2("2.1 Frontend")
add_bullets([
    "React 19, Vite 8 – Single Page Application",
    "TailwindCSS v4 – giao diện responsive, dark mode",
    "Framer Motion, GSAP, AOS, Swiper – hiệu ứng mượt",
    "Axios, React Router v7, Context API (Auth, Cart, Wishlist, Theme)",
    "Recharts – biểu đồ admin dashboard",
    "React Hot Toast – thông báo",
])
add_h2("2.2 Backend")
add_bullets([
    "Node.js, Express.js – kiến trúc MVC",
    "JWT + bcrypt – xác thực và phân quyền",
    "Multer – upload ảnh (sản phẩm, banner, trang, đánh giá)",
    "express-validator, helmet, rate-limit – bảo mật",
    "Nodemailer – email xác nhận đơn hàng (mock nếu chưa cấu hình SMTP)",
])
add_h2("2.3 Database")
add_bullets([
    "Microsoft SQL Server 2019+ / Express",
    "Database: MediCareStore",
    "Script cài đặt: SQL/SQLQuery1.sql (máy mới), SQL/install-safe.sql (cập nhật an toàn)",
    "Auto schema: backend/src/config/ensureSchema.js khi khởi động API",
])

add_h1("3. CẤU TRÚC THƯ MỤC")
add_bullets([
    "backend/ – API Express, uploads/, controllers, routes, middleware",
    "frontend/ – React SPA, components, pages, context, services",
    "SQL/ – SQLQuery1.sql (all-in-one), install-safe.sql",
    "database/ – schema.sql, seed.sql, migrations/",
    "CAI-DAT.bat + install-all.ps1 – cài đặt tự động 1 click",
    "note.txt, CAI-DAT-NHANH.txt – hướng dẫn cài đặt",
])

add_h1("4. CHỨC NĂNG HỆ THỐNG")
add_h2("4.1 Khách hàng (Storefront)")
add_bullets([
    "Trang chủ: banner slider, danh mục, sản phẩm nổi bật/bán chạy, FAQ, newsletter",
    "Danh sách sản phẩm: tìm kiếm, lọc danh mục/thương hiệu/giá, sắp xếp, phân trang",
    "Chi tiết sản phẩm: gallery zoom, thông số kỹ thuật, đánh giá có ảnh",
    "Giỏ hàng optimistic update, mã giảm giá, thanh toán COD/Online",
    "Yêu thích sản phẩm (Wishlist), xem nhanh (Quick View modal cố định giữa màn hình",
    "Nút Mua ngay đồng bộ trên thẻ sản phẩm, modal và trang chi tiết",
    "Đăng ký/đăng nhập JWT, quên mật khẩu, hồ sơ, lịch sử đơn hàng",
    "Trang Liên hệ (form + phản hồi admin), Giới thiệu (ảnh quản lý từ admin)",
    "Dark mode, header cố định, thông báo đơn hàng & liên hệ",
])
add_h2("4.2 Quản trị (Admin)")
add_bullets([
    "Dashboard: doanh thu, biểu đồ 6 tháng, thống kê đơn/sản phẩm/khách hàng",
    "CRUD sản phẩm: upload nhiều ảnh hoặc URL, ẩn/hiện, xóa vĩnh viễn",
    "CRUD danh mục: ẩn/hiện riêng, xóa vĩnh viễn",
    "Quản lý đơn hàng: trạng thái, phân trang",
    "Quản lý người dùng: khóa/mở, phân quyền, phân trang",
    "Kiểm duyệt đánh giá: duyệt/từ chối, xem ảnh đánh giá",
    "Quản lý liên hệ: trả lời khách, phân trang",
    "Cài đặt Banner trang chủ: upload/URL, ẩn/hiện",
    "Cài đặt ảnh trang Liên hệ & Giới thiệu",
])

add_h1("5. CƠ SỞ DỮ LIỆU – CÁC BẢNG CHÍNH")
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Bảng"
hdr[1].text = "Mô tả"
rows = [
    ("Users", "Tài khoản khách hàng và admin"),
    ("Products, ProductImages", "Sản phẩm và ảnh"),
    ("Categories, Brands", "Danh mục, thương hiệu"),
    ("Cart", "Giỏ hàng theo user"),
    ("Orders, OrderDetails, Payments", "Đơn hàng, chi tiết, thanh toán"),
    ("Coupons", "Mã giảm giá"),
    ("Reviews", "Đánh giá sản phẩm (có ImageUrls)"),
    ("Wishlists", "Sản phẩm yêu thích"),
    ("SiteSettings", "Cấu hình site, ảnh trang Liên hệ/Giới thiệu"),
    ("Contacts", "Form liên hệ + phản hồi admin"),
    ("Banners, Testimonials", "Banner trang chủ, feedback"),
    ("NewsletterSubscribers", "Đăng ký nhận tin"),
]
for a, b in rows:
    row = table.add_row().cells
    row[0].text = a
    row[1].text = b

add_h1("6. API ENDPOINTS CHÍNH")
add_h2("Public")
add_bullets([
    "GET /api/home – dữ liệu trang chủ",
    "GET /api/products, /api/products/:slug – sản phẩm",
    "GET /api/categories, /api/brands, /api/banners",
    "GET /api/pages/images – ảnh trang Liên hệ/Giới thiệu",
    "POST /api/auth/register, /api/auth/login",
])
add_h2("User (JWT)")
add_bullets([
    "GET/POST/DELETE /api/cart – giỏ hàng",
    "POST /api/orders – đặt hàng",
    "GET /api/orders/my – lịch sử đơn",
    "POST /api/reviews – đánh giá (upload ảnh)",
    "GET/POST/DELETE /api/wishlist – yêu thích",
    "POST /api/contacts – gửi liên hệ",
])
add_h2("Admin")
add_bullets([
    "GET /api/admin/dashboard – thống kê",
    "CRUD /api/products, /api/categories",
    "GET/PATCH /api/orders, /api/admin/users",
    "GET/PATCH /api/reviews, /api/contacts",
    "PUT /api/pages/images/:page/:slot – cập nhật ảnh trang",
])

add_h1("7. HƯỚNG DẪN CÀI ĐẶT")
add_h2("7.1 Yêu cầu")
add_bullets(["Node.js 18+", "SQL Server + SSMS", "npm"])
add_h2("7.2 Cài đặt nhanh (1 click)")
add_bullets([
    "Double-click file CAI-DAT.bat ở thư mục gốc project",
    "Lần đầu: sửa 3 dòng trong backend/.env: DB_SERVER, DB_USER, DB_PASSWORD",
    "Script tự động: npm install, chạy SQL, ensureSchema, tạo tài khoản demo",
])
add_h2("7.3 Chạy ứng dụng")
add_bullets([
    "Terminal 1: cd backend → npm run dev  (http://localhost:5000)",
    "Terminal 2: cd frontend → npm run dev (http://localhost:5173)",
    "Admin: http://localhost:5173/admin",
])
add_h2("7.4 Tài khoản demo")
add_bullets([
    "Admin: admin@medicarestore.com / Admin@123",
    "User: user@medicarestore.com / User@123",
    "Mã giảm giá: WELCOME10, MEDICARE50K, VIP15",
])
add_h2("7.5 Cấu hình SQL theo máy")
add_bullets([
    "SQL mặc định: DB_SERVER=localhost",
    "SQL Express: DB_SERVER=localhost\\SQLEXPRESS",
    "Tên máy Windows: DB_SERVER=TEN-MAY (vd: KAOOXANGG)",
    "DB_USER và DB_PASSWORD: tài khoản SQL Server Authentication",
])

add_h1("8. BẢO MẬT & XỬ LÝ LỖI")
add_bullets([
    "JWT token, bcrypt hash mật khẩu",
    "Helmet, CORS, rate limiting API",
    "Phân quyền admin/user qua middleware",
    "ensureSchema tự bổ sung bảng/cột thiếu khi khởi động backend",
    "Upload giới hạn kích thước, chỉ cho phép ảnh",
])

add_h1("9. GIAO DIỆN & TRẢI NGHIỆM NGƯỜI DÙNG")
add_bullets([
    "Responsive mobile/tablet/desktop",
    "Dark mode toàn site",
    "Hiệu ứng Framer Motion, hero slider fade",
    "Product card: trái tim yêu thích, icon mắt xem nhanh",
    "Buy Now gradient đồng bộ, optimistic cart",
    "Admin UI glass morphism, biểu đồ Recharts",
])

add_h1("10. KẾT LUẬN")
add_p(
    "Dự án MediCare Store hoàn thiện một hệ thống e-commerce thiết bị y tế fullstack "
    "với đầy đủ chức năng mua hàng, quản trị và mở rộng. Việc cài đặt được đơn giản hóa "
    "bằng CAI-DAT.bat và file SQL all-in-one SQLQuery1.sql, phù hợp triển khai trên máy "
    "mới hoặc môi trường demo/học tập."
)
add_p("")
add_p("Tài liệu tham khảo trong project: note.txt, CAI-DAT-NHANH.txt, README.md", bold=True)

doc.save(OUT)
print(f"Created: {OUT}")
